#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_word_doc():
    doc = Document()
    
    # 设置中文字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    title = doc.add_heading('二课活动管理系统 - 使用说明书', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    # 目录
    doc.add_heading('目录', 1)
    toc_items = [
        '系统简介',
        '学生使用指南',
        '活动负责人使用指南',
        '发布干事使用指南',
        '赋分干事使用指南',
        '管理员使用指南',
        '常见问题'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # 系统简介
    doc.add_heading('系统简介', 1)
    doc.add_paragraph('二课活动管理系统是一个用于管理第二课堂活动的平台，主要功能包括：')
    features = [
        '活动管理：发布、审核、管理各类第二课堂活动',
        '请假管理：学生提交请假申请，管理员审核',
        '活动赋分：对完成的活动进行学分赋分',
        '晚自习查询：查询晚自习请假记录'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_heading('访问地址', 2)
    doc.add_paragraph('系统网址：https://cf6e37ac-76fa-4ad2-b5a2-4ac4699268b4-5000.dev.coze.site')
    
    p = doc.add_paragraph()
    run = p.add_run('⚠️ 注意：如果网址无法访问，可能是系统休眠了。请联系管理员唤醒系统，或等待几分钟后重试。')
    run.font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_page_break()
    
    # 学生使用指南
    doc.add_heading('学生使用指南', 1)
    doc.add_heading('适用对象', 2)
    doc.add_paragraph('所有在校学生（无需注册账号）')
    
    doc.add_heading('功能一：提交请假申请', 2)
    doc.add_paragraph('使用场景：因事假、病假或参加公共活动需要请假时')
    doc.add_paragraph('操作步骤：')
    steps = [
        '打开系统首页',
        '点击「请假申请」卡片',
        '填写请假信息：学号、班级、姓名、请假类型、请假条图片（必填）、活动名称（活动公假时填写）',
        '点击「提交请假」',
        '提交成功后，可以查询请假状态'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('查询请假状态', 3)
    steps = [
        '点击首页「请假状态查询」',
        '输入学号或姓名',
        '查看审核状态（待审核 / 已通过 / 已驳回）'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('功能二：晚自习请假查询', 2)
    doc.add_paragraph('使用场景：查询晚自习请假记录')
    doc.add_paragraph('操作步骤：')
    steps = [
        '点击首页「晚自习请假查询」',
        '选择查询方式：按班级查询、按姓名查询、按学号查询',
        '查看查询结果'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_page_break()
    
    # 活动负责人使用指南
    doc.add_heading('活动负责人使用指南', 1)
    doc.add_heading('适用对象', 2)
    doc.add_paragraph('负责组织和申报活动的学生（需要注册账号）')
    
    doc.add_heading('第一步：注册账号', 2)
    steps = [
        '打开系统首页',
        '点击右上角「登录/注册」',
        '点击「去注册」',
        '填写信息：学号、姓名、密码、角色（选择「负责人」）',
        '点击「注册」'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('第二步：登录系统', 2)
    steps = [
        '点击首页「登录/注册」',
        '输入学号、姓名、密码',
        '点击「登录」',
        '登录成功后，首页会显示你的姓名和角色'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('功能一：提交活动', 2)
    doc.add_paragraph('使用场景：组织活动后，向学校申报活动信息')
    doc.add_paragraph('操作步骤：')
    steps = [
        '登录后，点击首页「活动提交」',
        '填写活动信息：活动名称、活动类别（德/智/体/美/劳）、活动级别（院系级/校级）、开始时间、结束时间、负责人姓名、负责人电话',
        '点击「提交活动」',
        '提交成功后，可以查询提交状态'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('查询提交状态', 3)
    steps = [
        '点击首页「提交状态查询」',
        '输入负责人手机号',
        '查看审核状态'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('功能二：提交赋分材料', 2)
    doc.add_paragraph('使用场景：活动结束后，提交赋分所需的材料')
    doc.add_paragraph('操作步骤：')
    steps = [
        '点击首页「赋分材料提交」',
        '输入负责人手机号查询已提交的活动',
        '选择要提交材料的活动',
        '上传材料：赋分表（必填）、备案表照片（校级活动需要）',
        '点击「提交材料」'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('功能三：修改密码', 2)
    steps = [
        '登录后，点击右上角你的姓名',
        '点击「修改密码」',
        '输入旧密码和新密码',
        '点击「确认修改」'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_page_break()
    
    # 发布干事使用指南
    doc.add_heading('发布干事使用指南', 1)
    doc.add_heading('适用对象', 2)
    doc.add_paragraph('负责审核活动发布的干事（需要管理员赋予权限）')
    
    doc.add_heading('登录系统', 2)
    steps = [
        '使用分配的账号登录（学号 + 姓名 + 密码）',
        '登录后，点击首页「发布活动」'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('功能：审核活动提交', 2)
    doc.add_paragraph('操作步骤：')
    steps = [
        '进入「发布活动」页面',
        '查看待审核的活动列表',
        '点击活动查看详细信息：活动基本信息、策划书（可下载查看）、备案表（可下载查看）',
        '选择审核结果：通过（活动审核通过，自动写入活动总表）或驳回（填写驳回原因，通知负责人）',
        '点击「确认」'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_page_break()
    
    # 赋分干事使用指南
    doc.add_heading('赋分干事使用指南', 1)
    doc.add_heading('适用对象', 2)
    doc.add_paragraph('负责对活动进行赋分的干事（需要管理员赋予权限）')
    
    doc.add_heading('登录系统', 2)
    steps = [
        '使用分配的账号登录',
        '登录后，点击首页「活动赋分」'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('功能：活动赋分', 2)
    doc.add_paragraph('操作步骤：')
    steps = [
        '进入「活动赋分」页面',
        '查看待赋分的活动列表',
        '点击活动查看详细信息：活动基本信息、赋分表（可下载查看）、备案表照片（校级活动，可下载查看）',
        '确认赋分材料无误后，点击「确认赋分」',
        '赋分完成后，系统会自动通知活动负责人'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('赋分规则', 3)
    doc.add_paragraph('院系级活动：只需审核赋分表')
    doc.add_paragraph('校级活动：需要审核赋分表 + 备案表照片')
    
    doc.add_page_break()
    
    # 管理员使用指南
    doc.add_heading('管理员使用指南', 1)
    doc.add_heading('适用对象', 2)
    doc.add_paragraph('系统管理员（拥有全部权限）')
    
    doc.add_heading('登录系统', 2)
    steps = [
        '使用管理员账号登录',
        '登录后，点击首页「管理员」'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('功能一：活动总表管理', 2)
    doc.add_heading('查看活动列表', 3)
    steps = [
        '进入管理后台，默认显示「活动总表」',
        '可以按类别、级别筛选活动',
        '可以搜索活动名称'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('添加活动', 3)
    steps = [
        '点击「添加活动」',
        '填写活动信息',
        '点击「确认添加」'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('编辑活动', 3)
    steps = [
        '点击活动右侧的「编辑」按钮',
        '修改活动信息',
        '点击「确认修改」'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('删除活动', 3)
    steps = [
        '点击活动右侧的「删除」按钮',
        '确认删除'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('功能二：活动审核', 2)
    steps = [
        '点击「活动审核」标签',
        '查看负责人提交的活动',
        '审核通过或驳回'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('功能三：请假审核', 2)
    steps = [
        '点击「请假审核」标签',
        '查看学生提交的请假申请',
        '查看请假条截图',
        '审核通过或驳回'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('功能四：活动赋分', 2)
    steps = [
        '点击「活动赋分」标签',
        '对待赋分的活动进行赋分操作'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('功能五：用户管理', 2)
    doc.add_heading('查看用户列表', 3)
    steps = [
        '点击「用户管理」标签',
        '查看所有注册用户'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('修改用户权限', 3)
    steps = [
        '找到要修改的用户',
        '点击「角色」下拉框，修改角色（学生/负责人/管理员）',
        '开启/关闭「发布活动权限」',
        '开启/关闭「活动赋分权限」',
        '开启/关闭「请假审核权限」'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('添加用户', 3)
    steps = [
        '点击「添加用户」',
        '填写用户信息',
        '点击「确认添加」'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_page_break()
    
    # 常见问题
    doc.add_heading('常见问题', 1)
    
    faqs = [
        ('Q1：网址打不开怎么办？', '原因：系统可能进入休眠状态\n解决方法：1. 等待 1-2 分钟后刷新页面 2. 联系管理员唤醒系统 3. 如果持续无法访问，可能是网络问题，稍后再试'),
        ('Q2：忘记密码怎么办？', '解决方法：1. 联系管理员重置密码 2. 或者重新注册一个新账号'),
        ('Q3：提交活动后多久能审核？', '审核时间由发布干事决定，一般 1-3 个工作日内完成审核'),
        ('Q4：请假申请被驳回了怎么办？', '1. 查看驳回原因 2. 修改后重新提交 3. 如有疑问，联系管理员'),
        ('Q5：如何知道我的活动是否赋分完成？', '1. 登录系统后，查看首页右上角的通知铃铛 2. 如果有新通知，铃铛上会显示红色数字 3. 点击铃铛查看通知详情'),
        ('Q6：活动 ID 是什么？', '活动 ID 格式：EK{年月}{序号}（如：EK202607001）\n活动 ID 由系统自动生成，不可修改\n用于唯一标识每个活动'),
        ('Q7：上传文件大小有限制吗？', '单个文件最大 10MB\n支持格式：图片（JPG、PNG）、PDF、Word 文档'),
        ('Q8：为什么我看不到某些功能入口？', '部分功能需要登录才能使用\n部分功能需要特定权限（如发布干事、赋分干事）\n联系管理员开通相应权限')
    ]
    
    for question, answer in faqs:
        doc.add_heading(question, 2)
        doc.add_paragraph(answer)
    
    doc.add_page_break()
    
    # 联系支持
    doc.add_heading('联系支持', 1)
    doc.add_paragraph('如有其他问题，请联系系统管理员：')
    doc.add_paragraph('管理员姓名：李广')
    doc.add_paragraph('学号：2505141139')
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('最后更新：2026 年 8 月')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 保存文档
    doc.save('/workspace/projects/二课活动管理系统使用说明书.docx')
    print('Word 文档已生成：/workspace/projects/二课活动管理系统使用说明书.docx')

if __name__ == '__main__':
    create_word_doc()
